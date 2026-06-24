from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Scaleway_Audit_Sentinel_Interview_Guide.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x18, 0x21, 0x2F)
MUTED = RGBColor(0x66, 0x70, 0x85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header_p = section.header.paragraphs[0]
    header_p.text = "Scaleway Audit Sentinel - Interview Guide"
    header_p.style = doc.styles["Normal"]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)


def add_title(doc):
    p = doc.add_paragraph(style="Title")
    p.add_run("Scaleway Audit Sentinel").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("Interview Study Guide: file-by-file explanation, runtime flows, design decisions, likely questions, and improvement ideas.")
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Positioning statement",
        "I chose the Scaleway assignment because it has a focused, defendable 2.5-hour scope: event ingestion, detection rules, dashboarding, remediation, auditability, Docker, and tests. The Coda assignment is broader because it spans documents, tables, pages, exports, content scanning, and several object-specific remediation paths.",
    )


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(text)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = True
        run.font.color.rgb = DARK_BLUE
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_geometry(table, widths, indent_dxa=120)
    doc.add_paragraph()
    return table


def add_flow_section(doc):
    doc.add_heading("Runtime Architecture", level=1)
    doc.add_paragraph(
        "The app is a standalone full-stack Node.js application. It serves a static browser dashboard and a small REST API from the same HTTP server."
    )
    add_table(
        doc,
        ["Layer", "Responsibility", "Key files"],
        [
            ["Browser dashboard", "Shows status, alerts, evidence, recent events, and remediation history.", "public/index.html, public/app.js, public/styles.css"],
            ["HTTP API", "Exposes status, alerts, events, remediation, manual polling, and static files.", "src/server.js, src/httpUtils.js"],
            ["Ingestion pipeline", "Fetches events on demand or on a schedule, then runs detection.", "src/poller.js"],
            ["Provider boundary", "Calls real Scaleway APIs or returns demo data using the same client shape.", "src/scalewayClient.js, src/demoData.js"],
            ["Detection engine", "Normalizes raw records and applies deterministic security rules.", "src/detectionRules.js"],
            ["Persistence and audit", "Stores events, alerts, principal IP history, user lock state, and remediation logs.", "src/store.js"],
        ],
        [1700, 4100, 3560],
    )
    doc.add_heading("End-to-End Flow", level=2)
    add_numbered(
        doc,
        [
            "Server starts and loads .env configuration.",
            "The store initializes data/state.json and the logger opens logs/app.log.",
            "The app chooses DemoScalewayClient for SCW_MODE=demo or ScalewayClient for SCW_MODE=live.",
            "The scheduler runs an initial scan and repeats based on SCW_POLL_INTERVAL_SECONDS.",
            "Audit and authentication events are fetched and normalized into one internal event model.",
            "Detection rules create alerts with severity, evidence, metadata, status, and remediation capabilities.",
            "The dashboard auto-refreshes and displays alerts, details, evidence, recent events, and remediation history.",
            "When an analyst clicks lock/unlock, the backend validates the alert and records the remediation action.",
        ],
    )


def add_file_sections(doc):
    doc.add_heading("File-by-File Explanation", level=1)

    root_rows = [
        ["package.json", "Project metadata and scripts. Declares Node >=20, ES modules, and scripts: start, test, poll.", "Mention that no external dependencies keeps setup and Docker fast."],
        [".env.example", "Runtime configuration template: mode, port, Scaleway credentials, polling interval, detection thresholds, country allowlist, data/log paths.", "Shows the app is configurable without source changes."],
        ["Dockerfile", "Builds a minimal Node 20 Alpine image, copies source/static files, exposes 3000, runs src/server.js.", "Satisfies Docker compatibility and reproducible setup."],
        [".gitignore", "Excludes generated and local-only files such as node_modules, data, logs, .env.", "Prevents credentials and runtime state from being committed."],
        ["README.md", "Submission guide with setup, architecture, live mode, API, rules, Docker, and AI usage.", "If asked why README exists despite the doc, it is for evaluator setup."],
    ]
    add_table(doc, ["File", "What it does", "Interview angle"], root_rows, [1900, 4850, 2610])

    backend_rows = [
        ["src/server.js", "Composition root. Loads config, initializes logger/store, chooses client, defines REST routes, serves frontend, starts scheduler.", "Main place to explain API design and dependency wiring."],
        ["src/config.js", "Loads .env, parses numbers/lists, builds a typed config object, validates required live-mode credentials.", "Fail-fast live configuration and environment-based setup."],
        ["src/scalewayClient.js", "Real REST client. Lists audit events/authentication events with pagination and calls IAM lock/unlock endpoints.", "Provider boundary; easy to mock, test, or replace."],
        ["src/demoData.js", "Credential-free client that returns sample audit/auth events and records lock state in memory.", "Makes the demo reliable without a Scaleway tenant."],
        ["src/detectionRules.js", "Normalizes raw events and runs rules: failed login burst, forbidden sensitive access, unusual country, new source IP, credential changes.", "Core security logic; easiest place to extend policies."],
        ["src/poller.js", "Fetches events, stores unique events, runs detection, upserts alerts, remembers principal IPs, updates poll metadata.", "Pipeline and scheduling behavior."],
        ["src/store.js", "File-backed JSON persistence for metadata, events, alerts, remediations, users, and principal IP history.", "Simple standalone storage; replaceable with DB in production."],
        ["src/remediation.js", "Validates alert/action, calls lock/unlock, updates local user state, changes alert status, records audit entry.", "Shows safe remediation workflow and auditability."],
        ["src/logger.js", "Writes JSON-line logs with levels to logs/app.log and console.", "Meets logging requirement and aids troubleshooting."],
        ["src/httpUtils.js", "Shared HTTP helpers for JSON parsing, JSON responses, error responses, and safe static file serving.", "Keeps server.js readable and avoids path traversal."],
        ["src/cli/poll-once.js", "Runs one detection cycle and prints JSON result.", "Useful for cron, debugging, or backend-only demo."],
    ]
    add_table(doc, ["File", "What it does", "Interview angle"], backend_rows, [1900, 4850, 2610])

    frontend_rows = [
        ["public/index.html", "Dashboard structure: header, analyst input, scan/refresh controls, metrics, alerts table, detail panel, events, remediation log.", "Shows the UI meets dashboard/remediation requirements."],
        ["public/app.js", "Frontend controller: fetches API data, renders alerts/details/events/remediations, handles filters, manual scan, lock/unlock, dismiss/reopen, auto-refresh.", "Vanilla JS keeps the demo dependency-free."],
        ["public/styles.css", "Responsive styling for metrics, panels, tables, badges, detail layout, mobile behavior.", "Professional but restrained operational UI."],
        ["public/logo.svg", "Small shield logo in header.", "Branding polish, not core logic."],
    ]
    add_table(doc, ["File", "What it does", "Interview angle"], frontend_rows, [1900, 4850, 2610])

    test_rows = [
        ["test/detectionRules.test.js", "Verifies demo events produce expected alert rule IDs and new-source-IP behavior requires existing profile history.", "Tests highest-risk detection logic."],
        ["test/store.test.js", "Verifies event deduplication and alert upsert/occurrence counting by fingerprint.", "Tests state management and duplicate prevention."],
    ]
    add_table(doc, ["File", "What it does", "Interview angle"], test_rows, [2200, 4750, 2410])


def add_detection_section(doc):
    doc.add_heading("Detection Rules in Detail", level=1)
    add_table(
        doc,
        ["Rule", "Severity", "Logic", "Why it matters"],
        [
            ["failed-login-burst", "High", "Counts failed authentication events per principal within FAILED_LOGIN_WINDOW_MINUTES and alerts when count reaches FAILED_LOGIN_THRESHOLD.", "Detects brute-force or credential-stuffing attempts."],
            ["forbidden-sensitive-access", "High", "Finds audit events with status 403 against sensitive hints such as IAM, secret, key, credential, token, password, or MFA.", "Flags privilege misuse or compromised accounts probing sensitive resources."],
            ["unusual-country", "Medium", "Raises alert when a successful authentication comes from a country not in ALLOWED_COUNTRY_CODES.", "Useful coarse anomaly check when geo source is available."],
            ["new-source-ip", "Medium", "Raises alert when a principal with known IP history successfully authenticates from an unseen source IP.", "Detects account access from a new network after a baseline exists."],
            ["credential-change", "Medium", "Detects successful lifecycle changes involving API keys, tokens, or MFA-related methods/resources.", "Credential creation/deletion is high-signal in cloud compromise investigations."],
        ],
        [1850, 1100, 4100, 2310],
    )
    doc.add_heading("Normalization Strategy", level=2)
    doc.add_paragraph(
        "Raw Scaleway audit records and authentication records do not have exactly the same shape. The app maps both into one internal event model with fields such as id, kind, recordedAt, actor, userId, sourceIp, countryCode, serviceName, methodName, statusCode, resourceTypes, resourceNames, metadata, and raw."
    )
    add_callout(
        doc,
        "Interview line",
        "I normalize first so the rule engine is independent of provider payload quirks. That makes rules easier to test and makes it realistic to support another cloud provider later.",
    )
    doc.add_heading("Alert Fingerprints", level=2)
    doc.add_paragraph(
        "Each alert gets a fingerprint such as failed-login-burst:user:hour or credential-change:event-id. The store upserts by fingerprint so repeated polling does not create endless duplicate rows. Existing alerts update lastSeenAt, occurrences, evidence, and metadata."
    )


def add_api_section(doc):
    doc.add_heading("API Surface", level=1)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/api/status", "GET", "Returns mode, region, polling status, last poll error, and count metrics."],
            ["/api/alerts?status=&severity=", "GET", "Returns alerts, optionally filtered by status and severity."],
            ["/api/events?limit=25", "GET", "Returns recent normalized events."],
            ["/api/remediations", "GET", "Returns local audit trail of remediation actions."],
            ["/api/poll", "POST", "Runs one ingestion/detection cycle on demand."],
            ["/api/alerts/:id/remediate", "POST", "Runs lock or unlock action and records remediation audit entry."],
            ["/api/alerts/:id/status", "PATCH", "Marks alert open, dismissed, or remediated."],
        ],
        [3300, 1050, 5010],
    )


def add_interview_section(doc):
    doc.add_heading("Likely Interview Questions and Strong Answers", level=1)
    qa = [
        ("Why did you pick Scaleway over Coda?", "Scaleway has a narrower, clearer 2.5-hour scope: event ingestion, detection, alert dashboard, remediation, logging, tests, Docker. Coda requires scanning documents, tables, rows, pages, exports, sharing state, and multiple remediation types, which is much broader."),
        ("Why no database?", "For a standalone coding assignment, JSON storage keeps setup simple and auditable. The store is isolated behind JsonStore, so replacing it with Postgres is straightforward when scale or concurrency requires it."),
        ("What happens if the Scaleway API fails?", "The poller catches the error, stores lastPollError, logs structured error data, and keeps the previous alerts/events available. In production I would add retries, backoff, and alerting on ingestion failure."),
        ("How do you avoid duplicate alerts?", "Events are deduplicated by event id. Alerts are upserted by fingerprint, so recurring detections update occurrences and evidence instead of creating noise."),
        ("How safe is remediation?", "The backend validates the alert exists, confirms it has a user target, validates the requested action, calls the provider client, updates local user lock state, changes alert status, and writes a remediation audit entry with actor, before, after, target, and timestamp."),
        ("Why demo mode?", "It makes the application reviewable without external credentials while keeping the same interface as the live client. That proves architecture without making the demo dependent on a tenant setup."),
        ("How would you scale this?", "Move state to Postgres, use a queue for ingestion jobs, run multiple workers with distributed locks, add API pagination, and separate the frontend from the API if traffic grows."),
        ("What security improvements would you make?", "Add analyst authentication/RBAC, CSRF protection if using cookies, immutable remediation audit logs, encrypted secrets, least-privilege Scaleway keys, and confirmation workflows for destructive actions."),
    ]
    for question, answer in qa:
        doc.add_heading(question, level=3)
        doc.add_paragraph(answer)


def add_improvements_section(doc):
    doc.add_heading("Changes and Improvements You Can Propose", level=1)
    add_table(
        doc,
        ["Area", "Improvement", "Why it is valuable"],
        [
            ["Security", "Add analyst login, RBAC, CSRF protection, and least-privilege API keys.", "Prevents unauthorized remediation actions."],
            ["Persistence", "Replace JsonStore with Postgres and indexes on recordedAt, ruleId, status, userId.", "Supports concurrent writes, history retention, and fast queries."],
            ["Reliability", "Add retry/backoff, rate-limit handling, and health checks.", "Makes ingestion resilient to temporary provider/API issues."],
            ["Detection", "Move rule configuration to YAML/JSON and add rule enable/disable controls.", "Allows policy changes without redeploying code."],
            ["Notifications", "Send Slack/email alerts for high severity issues.", "Meets optional alerting requirement and improves response time."],
            ["Remediation", "Add API key disable/rotate actions and MFA reset workflow.", "Responds more precisely than locking a user for every alert."],
            ["Auditability", "Make remediation log append-only and tamper-evident.", "Improves forensic trust and compliance posture."],
            ["Frontend", "Add pagination, server-side sorting, saved filters, and alert timeline view.", "Improves usability as event volume grows."],
            ["Testing", "Add integration tests with mocked Scaleway responses and UI smoke tests.", "Covers API contracts and user workflows."],
        ],
        [1650, 4100, 3610],
    )


def add_how_to_run_section(doc):
    doc.add_heading("How to Run and Demo", level=1)
    doc.add_heading("Local Demo Mode", level=2)
    add_numbered(
        doc,
        [
            "Copy .env.example to .env or rely on the default demo configuration.",
            "Run node --test to execute the unit tests.",
            "Run node src/server.js to start the app.",
            "Open http://127.0.0.1:3000.",
            "Click Run Scan and inspect generated alerts, evidence, recent events, and remediation log.",
        ],
    )
    doc.add_heading("Live Scaleway Mode", level=2)
    add_bullets(
        doc,
        [
            "Set SCW_MODE=live.",
            "Set SCW_SECRET_KEY, SCW_ORGANIZATION_ID, optional SCW_PROJECT_ID, and SCW_REGION.",
            "Use a key with the minimum permissions required to read audit/auth events and lock/unlock IAM users.",
            "Run node src/server.js and verify /api/status has no lastPollError.",
        ],
    )
    add_callout(
        doc,
        "Windows UNC note",
        "On this network path, npm.cmd may default to C:\\Windows. If that happens, run node --test and node src/server.js directly, or move/map the project to a drive letter.",
    )


def add_deep_dive_section(doc):
    doc.add_heading("Advanced Deep Dive Notes", level=1)
    doc.add_heading("Why the App Is Modular", level=2)
    add_bullets(
        doc,
        [
            "Provider integration is isolated in ScalewayClient and DemoScalewayClient.",
            "Detection logic is isolated in detectionRules.js and can be tested without HTTP or storage.",
            "Persistence is isolated in JsonStore, making a future database migration easier.",
            "Remediation orchestration is isolated in remediation.js so actions can be expanded safely.",
            "The UI only talks to REST endpoints, so the frontend could later move to React without changing backend internals.",
        ],
    )
    doc.add_heading("Failure Modes to Know", level=2)
    add_bullets(
        doc,
        [
            "If Scaleway API credentials are missing in live mode, startup fails fast.",
            "If a poll fails, previous dashboard data remains available and lastPollError is stored.",
            "If a remediation alert id is invalid, the API returns 404.",
            "If remediation is unsupported for an alert, the API returns 400.",
            "If JSON request parsing fails, the server logs the request failure and returns an error payload.",
        ],
    )
    doc.add_heading("Known Simplifications", level=2)
    add_bullets(
        doc,
        [
            "JsonStore is not ideal for concurrent multi-process writes.",
            "There is no analyst authentication in the demo.",
            "Geo/country detection relies on the event payload having country_code.",
            "The demo client does not persist lock state across process restarts.",
            "The live endpoint shapes are implemented from Scaleway docs, but a real tenant test should confirm exact field names and permissions.",
        ],
    )


def build():
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_flow_section(doc)
    add_file_sections(doc)
    add_detection_section(doc)
    add_api_section(doc)
    add_interview_section(doc)
    add_improvements_section(doc)
    add_how_to_run_section(doc)
    add_deep_dive_section(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
